"""Server-side .docx export of a BrandReport (python-docx).

Reproduces the on-screen report's section order and rating system in a
professionally styled Word document: navy headings, light-blue table header
rows, a color-coded grade banner and PASS/FAIL cells, the YoY table, monthly
trajectory, ranked red-flags table, category benchmarks, weighted scorecard,
the missing-data block under the grade, and the verdict.
"""

from __future__ import annotations

import io
from typing import Optional

from sales_support_agent.services.brand_analysis.schema import (
    CATEGORY_LABELS,
    BrandReport,
    benchmarks_for,
    fmt_money,
    fmt_mult,
    fmt_pct,
    safe_div,
)

# Brand palette (hex, no #) — mirrors the app design tokens.
_NAVY = "2B3644"
_LIGHT_BLUE = "85BBDA"
_HEADER_FILL = "DCE9F2"  # light-blue table header row
_MISSING_FILL = "F9F7F3"
_GRADE_FILL = {"A": "2E7D5B", "B": "3F8F6E", "C": "B8860B", "D": "C2663B", "F": "8B4C42"}
_SEV_HEX = {"Critical": "8B4C42", "High": "C2663B", "Medium": "B8860B"}
_PASS_HEX = "2E7D5B"
_FAIL_HEX = "8B4C42"


def _shade(cell, hex_fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _color_run(run, hex_color: str) -> None:
    from docx.shared import RGBColor

    run.font.color.rgb = RGBColor.from_string(hex_color)


def _heading(doc, text: str, level: int = 1) -> None:
    from docx.shared import RGBColor

    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(_NAVY)


def _kv_table(doc, headers: list[str], rows: list[list], *, align_right_from: int = 1) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        _shade(hdr[i], _HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
            if i >= align_right_from:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def build_docx(report: BrandReport) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    r = report
    doc = Document()

    # Header
    title = doc.add_paragraph()
    run = title.add_run("EXECUTIVE ACQUISITION REPORT")
    run.bold = True
    run.font.size = Pt(11)
    _color_run(run, _LIGHT_BLUE)
    _heading(doc, r.brand or "Brand", level=0)
    detected = ", ".join(r.detected_brands) if r.detected_brands else "—"
    periods = r.period_current_label + (f" vs {r.period_prior_label}" if r.has_yoy else " (single period — no prior year)")
    meta = doc.add_paragraph()
    meta.add_run(
        f"Category: {CATEGORY_LABELS.get(r.category, r.category)}  ·  Detected: {detected}  ·  "
        f"Periods: {periods}" + (f"  ·  Prepared {r.prepared_date}" if r.prepared_date else "")
    ).italic = True

    # 0. Grade banner
    banner = doc.add_table(rows=1, cols=2)
    cell_grade, cell_text = banner.rows[0].cells
    cell_grade.text = ""
    g_run = cell_grade.paragraphs[0].add_run(r.scorecard.letter)
    g_run.bold = True
    g_run.font.size = Pt(36)
    g_run.font.color.rgb = RGBColor.from_string("FFFFFF")
    cell_grade.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _shade(cell_grade, _GRADE_FILL.get(r.scorecard.letter, _NAVY))
    cell_text.paragraphs[0].add_run(
        f"Weighted score {r.scorecard.score_100}/100 · Recommendation: {r.recommendation}"
    ).bold = True
    cell_text.add_paragraph(r.verdict_text)
    doc.add_paragraph()

    # Missing-data block directly under the grade
    mp = doc.add_paragraph()
    if r.data_sufficient:
        run = mp.add_run(f"Data sufficient — grade reflects complete financial inputs.  (Confidence: {r.confidence})")
        run.bold = True
        _color_run(run, _PASS_HEX)
    else:
        run = mp.add_run(f"Missing data that would raise confidence  (Confidence: {r.confidence})")
        run.bold = True
        for item in r.missing_data:
            doc.add_paragraph(item, style="List Bullet")

    # 1. Executive Summary
    _heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(r.executive_summary)
    doc.add_paragraph().add_run("What stands out beyond standard KPIs").bold = True
    for s in r.stands_out:
        doc.add_paragraph(s, style="List Bullet")

    # 2. Financial Overview (YoY)
    _heading(doc, "2. Financial Overview (YoY)", 1)
    c, p = r.current, r.prior
    headers = ["Metric", r.period_current_label] + ([r.period_prior_label] if r.has_yoy else [])

    def _row(label, cur, prior):
        return [label, cur] + ([prior] if r.has_yoy else [])

    rows = [
        _row("Net revenue", fmt_money(c.net_revenue_cents), fmt_money(p.net_revenue_cents)),
        _row("COGS", fmt_money(c.cogs_cents), fmt_money(p.cogs_cents)),
        _row("Product gross profit", fmt_money(c.product_gross_profit_cents), fmt_money(p.product_gross_profit_cents)),
        _row("Product GM%", fmt_pct(c.product_gm_bps), fmt_pct(p.product_gm_bps)),
        _row("Marketing spend", fmt_money(c.marketing_total_cents), fmt_money(p.marketing_total_cents)),
        _row("Marketing % of revenue", fmt_pct(c.marketing_pct_bps), fmt_pct(p.marketing_pct_bps)),
        _row("Blended MER", fmt_mult(c.blended_mer), fmt_mult(p.blended_mer)),
        _row("Contribution (reported GP)", fmt_money(c.reported_gross_profit_cents), fmt_money(p.reported_gross_profit_cents)),
        _row("Contribution margin", fmt_pct(c.contribution_margin_bps), fmt_pct(p.contribution_margin_bps)),
        _row("Operating expenses", fmt_money(c.opex_cents), fmt_money(p.opex_cents)),
        _row("Net earnings", fmt_money(c.net_earnings_cents), fmt_money(p.net_earnings_cents)),
        _row("Net margin", fmt_pct(c.net_margin_bps), fmt_pct(p.net_margin_bps)),
    ]
    _kv_table(doc, headers, rows)
    if r.has_yoy:
        growth = "—" if r.yoy_revenue_growth_bps is None else fmt_pct(r.yoy_revenue_growth_bps)
        doc.add_paragraph(f"YoY net-revenue growth: {growth}.")

    # Monthly trajectory
    if r.monthly_revenue:
        doc.add_paragraph().add_run("Monthly revenue trajectory").bold = True
        peak = max((v for _, v in r.monthly_revenue if isinstance(v, (int, float))), default=1) or 1
        mt = doc.add_table(rows=0, cols=3)
        mt.style = "Light List Accent 1"
        for lbl, v in r.monthly_revenue:
            cells = mt.add_row().cells
            cells[0].text = str(lbl)
            cells[1].text = fmt_money(v)
            blocks = max(1, round((v / peak) * 20)) if isinstance(v, (int, float)) else 0
            cells[2].text = "█" * blocks
        doc.add_paragraph()

    # 3. Acquisition Evaluation
    _heading(doc, "3. Acquisition Evaluation", 1)
    bm = benchmarks_for(r.category)
    acq_cur = r.acquisition_current
    acq_pri = r.acquisition_prior

    # Grade badge + rationale
    acq_dim = next((d for d in r.scorecard.dimensions if d.key == "acquisition"), None)
    if acq_dim:
        badge = doc.add_table(rows=1, cols=2)
        badge.style = "Light Grid Accent 1"
        gcell, rcell = badge.rows[0].cells
        gcell.text = ""
        gr = gcell.paragraphs[0].add_run(acq_dim.letter)
        gr.bold = True
        gr.font.size = Pt(28)
        gr.font.color.rgb = RGBColor.from_string(_GRADE_FILL.get(acq_dim.letter, _NAVY))
        gcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(gcell, _MISSING_FILL)
        rcell.text = ""
        rcell.paragraphs[0].add_run("Acquisition mix & dependency  ·  12% weight").bold = True
        rcell.add_paragraph(acq_dim.reason)
        doc.add_paragraph()

    # Helper: write a PASS/FAIL/gap cell
    def _vd(cell, passed):
        cell.text = ""
        if passed is True:
            run = cell.paragraphs[0].add_run("PASS"); run.bold = True; _color_run(run, _PASS_HEX)
        elif passed is False:
            run = cell.paragraphs[0].add_run("FAIL"); run.bold = True; _color_run(run, _FAIL_HEX)
        else:
            cell.text = "—"

    # ── 3a. Customer Revenue Split ──────────────────────────────────────────
    doc.add_paragraph().add_run("Customer Revenue Split").bold = True
    new_r = acq_cur.get("new_customer_revenue_cents")
    ret_r = acq_cur.get("returning_customer_revenue_cents")
    new_r_p = acq_pri.get("new_customer_revenue_cents")
    ret_r_p = acq_pri.get("returning_customer_revenue_cents")
    coh_total = (new_r or 0) + (ret_r or 0) if (new_r is not None or ret_r is not None) else None
    coh_total_p = (new_r_p or 0) + (ret_r_p or 0) if (new_r_p is not None or ret_r_p is not None) else None
    ret_pct_bps = round(ret_r / coh_total * 10000) if (ret_r is not None and coh_total) else None
    ret_pct_bps_p = round(ret_r_p / coh_total_p * 10000) if (ret_r_p is not None and coh_total_p) else None
    new_pct_bps = round(new_r / coh_total * 10000) if (new_r is not None and coh_total) else None
    new_pct_bps_p = round(new_r_p / coh_total_p * 10000) if (new_r_p is not None and coh_total_p) else None

    def _pr(v):
        return "—" if v is None else v

    cols = ["Metric", r.period_current_label] + ([r.period_prior_label] if r.has_yoy else []) + ["Healthy", "Verdict"]
    split_table = doc.add_table(rows=1, cols=len(cols))
    split_table.style = "Light Grid Accent 1"
    hdr = split_table.rows[0].cells
    for i, h in enumerate(cols):
        hdr[i].text = ""
        hdr[i].paragraphs[0].add_run(h).bold = True
        _shade(hdr[i], _HEADER_FILL)

    def _split_row(label, cur_val, pri_val, healthy, passed):
        cells = split_table.add_row().cells
        cells[0].text = label
        cells[1].text = _pr(cur_val)
        idx = 2
        if r.has_yoy:
            cells[idx].text = _pr(pri_val); idx += 1
        cells[idx].text = healthy; idx += 1
        _vd(cells[idx], passed)

    _split_row("New-customer revenue", fmt_money(new_r), fmt_money(new_r_p), "—", None)
    _split_row("Returning-customer revenue", fmt_money(ret_r), fmt_money(ret_r_p), "—", None)
    _split_row("Returning-customer share",
               fmt_pct(ret_pct_bps), fmt_pct(ret_pct_bps_p),
               "≥ 30%",
               None if ret_pct_bps is None else ret_pct_bps >= 3000)
    _split_row("New-customer share",
               fmt_pct(new_pct_bps), fmt_pct(new_pct_bps_p),
               "—", None)
    _split_row("Owned-channel (email/SMS) %",
               fmt_pct(c.owned_pct_bps), fmt_pct(p.owned_pct_bps),
               f"{bm.owned_pct_bps[0]//100}–{bm.owned_pct_bps[1]//100}%",
               None if c.owned_pct_bps is None else c.owned_pct_bps >= bm.owned_pct_bps[0])
    doc.add_paragraph()

    # ── 3b. Retention & Pricing Signals ────────────────────────────────────
    doc.add_paragraph().add_run("Retention & Pricing Signals").bold = True
    sig_table = doc.add_table(rows=1, cols=len(cols))
    sig_table.style = "Light Grid Accent 1"
    hdr2 = sig_table.rows[0].cells
    for i, h in enumerate(cols):
        hdr2[i].text = ""
        hdr2[i].paragraphs[0].add_run(h).bold = True
        _shade(hdr2[i], _HEADER_FILL)

    def _sig_row(label, cur_val, pri_val, healthy, passed):
        cells = sig_table.add_row().cells
        cells[0].text = label
        cells[1].text = _pr(cur_val)
        idx = 2
        if r.has_yoy:
            cells[idx].text = _pr(pri_val); idx += 1
        cells[idx].text = healthy; idx += 1
        _vd(cells[idx], passed)

    _sig_row("Discount rate",
             fmt_pct(c.discount_rate_bps), fmt_pct(p.discount_rate_bps),
             f"{bm.discount_rate_bps[0]//100}–{bm.discount_rate_bps[1]//100}%",
             None if c.discount_rate_bps is None else c.discount_rate_bps <= bm.discount_rate_bps[1])
    _sig_row("Return rate",
             fmt_pct(c.return_rate_bps), fmt_pct(p.return_rate_bps),
             f"< {bm.return_rate_max_bps//100}%",
             None if c.return_rate_bps is None else c.return_rate_bps < bm.return_rate_max_bps)
    _sig_row("Blended MER",
             fmt_mult(c.blended_mer), fmt_mult(p.blended_mer),
             f"≥ {bm.blended_mer_min:.1f}x",
             None if c.blended_mer is None else c.blended_mer >= bm.blended_mer_min)
    _sig_row("Marketing % of revenue",
             fmt_pct(c.marketing_pct_bps), fmt_pct(p.marketing_pct_bps),
             f"{bm.marketing_pct_bps[0]//100}–{bm.marketing_pct_bps[1]//100}%",
             None if c.marketing_pct_bps is None else
             bm.marketing_pct_bps[0] <= c.marketing_pct_bps <= bm.marketing_pct_bps[1])
    doc.add_paragraph()

    # ── 3c. Unit Economics ──────────────────────────────────────────────────
    aov = acq_cur.get("aov_cents")
    cac = acq_cur.get("cac_cents")
    ltv = acq_cur.get("ltv_cents")
    ltv_cac = safe_div(ltv, cac) if (ltv is not None and cac) else None
    has_unit = any(v is not None for v in (aov, cac, ltv))
    doc.add_paragraph().add_run("Unit Economics").bold = True
    if has_unit or not coh_total:
        ue_rows = [
            ["AOV (average order value)", fmt_money(aov) if aov is not None else "Data gap — not supplied"],
            ["CAC (customer acquisition cost)", fmt_money(cac) if cac is not None else "Data gap — not supplied"],
            ["LTV (customer lifetime value)", fmt_money(ltv) if ltv is not None else "Data gap — not supplied"],
            ["LTV : CAC ratio", (f"{ltv_cac:.1f}x" + (" ✓" if ltv_cac and ltv_cac >= 3 else " (healthy ≥ 3x)")) if ltv_cac is not None else "Data gap — need LTV and CAC"],
        ]
        _kv_table(doc, ["Metric", "Value"], ue_rows, align_right_from=99)
    else:
        doc.add_paragraph("AOV, CAC, and LTV not supplied — request cohort export or platform-level CAC report.")

    # 4. Media Mix
    _heading(doc, "4. Media Mix", 1)
    if r.media_mix:
        total = sum(r.media_mix.values()) or 1
        mrows = [[k, fmt_money(v), f"{v/total*100:.0f}%"] for k, v in sorted(r.media_mix.items(), key=lambda kv: kv[1], reverse=True)]
        _kv_table(doc, ["Channel", "Spend", "% allocation"], mrows, align_right_from=1)
    else:
        doc.add_paragraph("Channel-level media mix not supplied — request ad-platform exports (Meta, Google, TikTok).")

    # 5. Contribution & Unit Economics
    _heading(doc, "5. Contribution & Unit Economics", 1)
    _kv_table(doc, ["Metric", "Value"], [
        ["Contribution margin", fmt_pct(c.contribution_margin_bps)],
        ["Discount rate", fmt_pct(c.discount_rate_bps)],
        ["Operating result excl. other income", fmt_money(c.operating_result_ex_other_cents)],
        ["Net earnings (reported)", fmt_money(c.net_earnings_cents)],
        ["AOV / CAC / LTV / payback", "Data gap — not supplied"],
    ], align_right_from=99)

    # 6. Balance Sheet & Earnings Quality
    _heading(doc, "6. Balance Sheet & Earnings Quality", 1)
    if r.balance_sheet:
        _kv_table(doc, ["Line", "Amount"], [[l, fmt_money(v)] for l, v in r.balance_sheet], align_right_from=1)
        if r.related_party_flag:
            warn = doc.add_paragraph().add_run("Related-party / intercompany items detected — diligence collectability and agreements.")
            _color_run(warn, _FAIL_HEX)
            warn.bold = True
    else:
        doc.add_paragraph("Balance sheet not supplied — request assets, intercompany balances, equity, dividends and related-party agreements.")

    # 7. Red Flags
    _heading(doc, "7. Red Flags", 1)
    if r.red_flags:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Severity", "Finding"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
            _shade(hdr[i], _HEADER_FILL)
        for f in r.red_flags:
            cells = table.add_row().cells
            sev_run = cells[0].paragraphs[0].add_run(f.severity)
            sev_run.bold = True
            _color_run(sev_run, _SEV_HEX.get(f.severity, _NAVY))
            cells[1].paragraphs[0].add_run(f.title).bold = True
            if f.detail:
                cells[1].add_paragraph(f.detail)
        doc.add_paragraph()
    else:
        doc.add_paragraph("No material red flags surfaced in the supplied data.")

    # 8. Category Benchmarks
    _heading(doc, "8. Category Benchmarks", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["KPI", "Healthy range", "Brand actual", "Verdict"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        _shade(hdr[i], _HEADER_FILL)
    for b in r.benchmarks:
        cells = table.add_row().cells
        cells[0].text = b.kpi
        cells[1].text = b.healthy
        cells[2].text = b.actual
        if b.passed is True:
            vr = cells[3].paragraphs[0].add_run("PASS"); vr.bold = True; _color_run(vr, _PASS_HEX)
        elif b.passed is False:
            vr = cells[3].paragraphs[0].add_run("FAIL"); vr.bold = True; _color_run(vr, _FAIL_HEX)
        else:
            cells[3].text = "data gap"
    doc.add_paragraph()

    # Weighted Scorecard
    _heading(doc, "Weighted Scorecard", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Dimension", "Weight", "Grade", "Reason"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        _shade(hdr[i], _HEADER_FILL)
    for d in r.scorecard.dimensions:
        cells = table.add_row().cells
        cells[0].text = d.label
        cells[1].text = f"{int(d.weight*100)}%"
        gr = cells[2].paragraphs[0].add_run(d.letter); gr.bold = True; _color_run(gr, _GRADE_FILL.get(d.letter, _NAVY))
        cells[3].text = d.reason
    doc.add_paragraph(f"Weighted composite (A=4…F=0, rebased to 100): {r.scorecard.score_100}/100 → {r.scorecard.letter}.")

    # 9. Data Gaps to Close
    _heading(doc, "9. Data Gaps to Close", 1)
    for g in r.data_gaps:
        doc.add_paragraph(g, style="List Bullet")

    # 10. Verdict
    _heading(doc, "10. Verdict", 1)
    vp = doc.add_paragraph()
    vp.add_run(f"Grade {r.scorecard.letter} ({r.scorecard.score_100}/100) — {r.recommendation}. ").bold = True
    vp.add_run(r.verdict_text)
    note = doc.add_paragraph()
    note.add_run(
        f"Source: derived from uploaded financial statements. Narrative: {r.narrative_model}. {r.intake_summary}"
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
