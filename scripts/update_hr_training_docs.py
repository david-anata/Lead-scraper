"""Refresh the HR training guides after the payroll-control-room hardening."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OPERATOR = ROOT / "docs/training/Anata-HR-Operator-Training-Guide.docx"
EMPLOYEE = ROOT / "docs/training/Anata-Employee-HR-Quick-Start.docx"


def replace(document: Document, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text == old:
            paragraph.text = new
            return
    raise ValueError(f"Paragraph not found: {old}")


def insert_before(document: Document, anchor: str, items: list[tuple[str, str]]) -> None:
    target = next(p for p in document.paragraphs if p.text == anchor)
    for text, style in items:
        paragraph = document.add_paragraph(text, style=style)
        target._p.addprevious(paragraph._p)


def update_operator() -> None:
    document = Document(OPERATOR)
    replace(
        document,
        "A qualified payroll or tax professional reviews the 2026 setup and opening balances. Mark the review complete only after that review actually occurs.",
        "A qualified payroll or tax professional reviews the 2026 setup and opening balances. Record the reviewer’s name, email, review date, evidence reference, note, and attestation. A checkbox alone does not complete this control.",
    )
    replace(
        document,
        "Finish employee setup, W-4s, open punches, time corrections, pending PTO, opening balances, tax review, and pending payroll inputs.",
        "Finish employee setup, W-4s, open punches, time corrections, employee-submitted timesheets, independent timesheet approval, pending PTO, independently approved opening balances, qualified tax review, and pending payroll inputs.",
    )
    replace(
        document,
        "Until bank transfer is connected, fill and deliver paper checks using Anata’s calculated net amounts. Confirm check number and issuance in the run.",
        "Until a full-service payroll provider is selected, fill and deliver paper checks using the approved amounts. Confirm each check number and issuance in the run. Do not describe an Anata estimate as provider-authoritative payroll.",
    )
    replace(
        document,
        "Reports are permission-filtered CSV exports. Store exports only where authorized HR administrators can access them.",
        "Reports are permission-filtered CSV exports. Quarterly and year-to-date registers support accountant reconciliation. The verified HR backup ZIP includes a checksum manifest but excludes full SSNs and sealed tax-form payloads. Store every download securely.",
    )
    insert_before(
        document,
        "9. Contractor payments through Wise",
        [
            ("9. Outside payroll-provider handoff", "Heading 1"),
            ("1. Export the approved run", "Heading 2"),
            ("Open the approved payroll version and download the provider handoff CSV. It contains hours, approved variable inputs, Anata estimates, and snapshot hashes; it does not transmit money.", "Normal"),
            ("2. Enter payroll in the outside provider", "Heading 2"),
            ("The outside provider must perform the legally authoritative wage calculation, tax withholding/deposit/filing, and direct deposit when that service is connected.", "Normal"),
            ("3. Record the provider run", "Heading 2"),
            ("Enter the provider name, provider run/reference, and an evidence note. Recording submission does not mean payroll has been paid or taxes have been filed.", "Normal"),
            ("4. Compare final totals", "Heading 2"),
            ("Enter the provider’s final gross, net, total taxes, and employer cost. Anata marks an exact match or lists each variance. Investigate every variance before relying on the outside result.", "Normal"),
            ("5. Preserve the evidence", "Heading 2"),
            ("Keep the final provider report and confirmation under Anata’s secure records policy. Anata keeps the reference and comparison audit; the provider remains the authority for its filings and transfers.", "Normal"),
        ],
    )
    # Renumber the remaining top-level sections after inserting the provider chapter.
    for old, new in (
        ("9. Contractor payments through Wise", "10. Contractor payments through Wise"),
        ("10. Pay statements, reports, and policies", "11. Pay statements, reports, and policies"),
        ("11. Offboarding", "12. Offboarding"),
        ("12. Troubleshooting quick reference", "13. Troubleshooting quick reference"),
        ("13. Payroll-day checklist", "14. Payroll-day checklist"),
    ):
        replace(document, old, new)
    insert_before(
        document,
        "13. Troubleshooting quick reference",
        [
            ("Compliance calendar and reminders", "Heading 2"),
            ("Open HR → Compliance to review all 24 paydays, Utah new-hire deadlines, quarterly Form 941, Utah withholding and unemployment reports, FUTA, W-2/W-3, and annual Utah reconciliation. Record an item complete only with the outside confirmation and evidence note.", "Normal"),
            ("Mobile employee use", "Heading 2"),
            ("On a phone, employees use the bottom shortcuts for Home, Time, Pay, and Profile. Dense records scroll inside their table instead of moving the entire page sideways.", "Normal"),
        ],
    )
    insert_before(
        document,
        "Correct period and pay date.",
        [
            ("Employee timesheets submitted and independently approved.", "List Bullet"),
            ("Opening balances independently approved and unchanged.", "List Bullet"),
            ("Compliance deadlines reviewed for the pay date.", "List Bullet"),
            ("Outside provider reference and final-total comparison recorded when a provider is used.", "List Bullet"),
        ],
    )
    document.save(OPERATOR)


def update_employee() -> None:
    document = Document(EMPLOYEE)
    insert_before(
        document,
        "3. Request PTO",
        [
            ("3. Submit your timesheet", "Heading 1"),
            ("1. Review the pay period", "Heading 2"),
            ("Confirm every workday is present, every punch is closed, and any correction request has been resolved.", "Normal"),
            ("2. Sign and submit", "Heading 2"),
            ("Confirm the timesheet is complete and accurate, then submit it. Another authorized person must approve it before payroll can be prepared.", "Normal"),
            ("3. Resubmit after changes", "Heading 2"),
            ("If approved time later changes, the approval becomes stale. Review the corrected time and submit the period again.", "Normal"),
        ],
    )
    for old, new in (
        ("3. Request PTO", "4. Request PTO"),
        ("4. View your pay statement", "5. View your pay statement"),
    ):
        replace(document, old, new)
    insert_before(
        document,
        "Need help?",
        [
            ("6. Use HR on your phone", "Heading 1"),
            ("Use the bottom shortcuts: Home for open tasks, Time for clocking and PTO, Pay for issued statements, and Profile for onboarding and secure forms.", "Normal"),
            ("Tables scroll sideways inside the record when needed. Your full Social Security number and sealed tax form are never displayed in those tables or downloads.", "Normal"),
        ],
    )
    document.save(EMPLOYEE)


if __name__ == "__main__":
    update_operator()
    update_employee()
